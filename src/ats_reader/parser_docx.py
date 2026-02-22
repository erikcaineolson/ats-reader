"""DOCX resume extraction via python-docx."""

from __future__ import annotations

from pathlib import Path
from xml.etree.ElementTree import Element

from docx import Document
from docx.oxml.ns import qn

from ats_reader.models import ExtractionMetadata


# Style info attached to each paragraph for structure analysis
class ParagraphStyle:
    """Lightweight record of a paragraph's formatting."""

    def __init__(
        self,
        text: str,
        style_name: str | None,
        is_heading: bool,
        heading_level: int | None,
        is_bold: bool,
        font_size: float | None,
        is_list_style: bool,
        has_manual_bullet: bool,
    ):
        self.text = text
        self.style_name = style_name
        self.is_heading = is_heading
        self.heading_level = heading_level
        self.is_bold = is_bold
        self.font_size = font_size
        self.is_list_style = is_list_style
        self.has_manual_bullet = has_manual_bullet


def parse_docx(path: Path) -> tuple[str, ExtractionMetadata, list[ParagraphStyle]]:
    """Extract text, metadata, and paragraph style info from a DOCX resume.

    Returns (raw_text, metadata, paragraph_styles).
    """
    doc = Document(str(path))
    meta = ExtractionMetadata(file_type="docx", page_count=_estimate_pages(doc))

    paragraphs_text: list[str] = []
    paragraph_styles: list[ParagraphStyle] = []
    fonts: set[str] = set()

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            paragraphs_text.append("")
            continue
        paragraphs_text.append(text)

        # Style analysis
        style_name = para.style.name if para.style else None
        is_heading = style_name is not None and style_name.startswith("Heading")
        heading_level = None
        if is_heading and style_name:
            try:
                heading_level = int(style_name.split()[-1])
            except (ValueError, IndexError):
                pass

        is_bold = _para_is_bold(para)
        font_size = _para_font_size(para)
        is_list_style = style_name in ("List Paragraph", "List Bullet", "List Number")
        has_manual_bullet = _has_manual_bullet(text)

        paragraph_styles.append(
            ParagraphStyle(
                text=text,
                style_name=style_name,
                is_heading=is_heading,
                heading_level=heading_level,
                is_bold=is_bold,
                font_size=font_size,
                is_list_style=is_list_style,
                has_manual_bullet=has_manual_bullet,
            )
        )

        # Collect fonts
        for run in para.runs:
            if run.font and run.font.name:
                fonts.add(run.font.name)

    # Tables
    if doc.tables:
        meta.has_tables = True
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs_text.append("  |  ".join(cells))

    # Images (inline shapes)
    meta.has_images = _has_images(doc)

    # Text boxes
    meta.has_text_boxes = _has_text_boxes(doc)

    # Headers and footers
    h_f = _extract_headers_footers(doc)
    if h_f:
        meta.has_headers_footers = True
        meta.header_footer_text = h_f

    meta.fonts_used = sorted(fonts)

    raw_text = "\n".join(paragraphs_text)
    return raw_text, meta, paragraph_styles


def _estimate_pages(doc: Document) -> int:
    """Rough page estimate based on paragraph count."""
    # DOCX doesn't have a reliable page count without rendering.
    # Rough heuristic: ~40 paragraphs per page.
    count = len(doc.paragraphs)
    return max(1, (count + 39) // 40)


def _para_is_bold(para) -> bool:
    """Check if the entire paragraph is bold."""
    if not para.runs:
        return False
    return all(run.bold for run in para.runs if run.text.strip())


def _para_font_size(para) -> float | None:
    """Get the font size (in pt) of the first run with text."""
    for run in para.runs:
        if run.text.strip() and run.font and run.font.size:
            return run.font.size.pt
    return None


MANUAL_BULLETS = {"•", "●", "○", "■", "◆", "▪", "–", "—", "-", "►", "➢", "✓", "✔"}


def _has_manual_bullet(text: str) -> bool:
    """Check if the line starts with a manual bullet character."""
    if not text:
        return False
    return text[0] in MANUAL_BULLETS


def _has_images(doc: Document) -> bool:
    """Check for inline images in the document body."""
    for para in doc.paragraphs:
        for run in para.runs:
            drawing_elements = run._element.findall(qn("w:drawing"))
            if drawing_elements:
                return True
    return False


def _has_text_boxes(doc: Document) -> bool:
    """Check for text boxes (w:txbxContent) in the document XML."""
    body: Element = doc.element.body
    txbx = body.findall(".//" + qn("w:txbxContent"))
    return len(txbx) > 0


def _extract_headers_footers(doc: Document) -> list[str]:
    """Extract text from document headers and footers."""
    texts: list[str] = []
    for section in doc.sections:
        if section.header and section.header.paragraphs:
            for para in section.header.paragraphs:
                t = para.text.strip()
                if t:
                    texts.append(t)
        if section.footer and section.footer.paragraphs:
            for para in section.footer.paragraphs:
                t = para.text.strip()
                if t:
                    texts.append(t)
    return texts
